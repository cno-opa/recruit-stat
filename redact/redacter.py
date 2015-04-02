#!/usr/bin/python

import re
import glob
import docx
from docx import Document

docs = glob.glob('*.docx')

# find info to be redacted
def find_vitals(paragraphs):

    global ifname
    global ilname
    global fname
    global lname
    global address
    global ssn

    paragraphs = paragraphs

    # limit size of doc to search to text above second line of asterisks
    lower_limit = None
    for p in paragraphs:
        if re.match(r'\*\*\*\*', p.text):
            lower_limit = paragraphs.index(p)

    paragraphs[0:lower_limit]

    for p in paragraphs:
        t = p.text.lower()

        if re.search('investigator', t):
            parts = re.split('\\t| ', t)
            ifname = parts[2]
            ilname = parts[3]

        if re.search('name:', t):
            parts = re.split('\\t| ', t)
            fname = parts[1]
            lname = parts[2]

        if re.search('address:', t):
            parts = re.split('\\t| ', t)
            address = parts[1]

        if re.search('s.s.n.:', t):
            parts = re.split('\\t| ', t)
            ssn = parts[1]

# redact vitals
def redact(redacters, doc):
    redacted = Document()

    for p in doc.paragraphs:
        p_r = p.text.lower()
        for r in redacters:
            p_r = re.sub(r, '[REDACTED]', p_r)
        redacted.add_paragraph(p_r)
    return redacted

# run
for doc in docs:
    doc_obj = Document(doc)
    find_vitals(doc_obj.paragraphs)
    redacters = [ifname, ilname, fname, lname, address, ssn]
    redacted = redact(redacters, doc_obj)

    redacted.save(doc + '-REDACTED.docx')
